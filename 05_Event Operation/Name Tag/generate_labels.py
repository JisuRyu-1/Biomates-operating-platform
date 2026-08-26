from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from openpyxl import load_workbook


TEMPLATE_PATH = "9106.docx"
EXCEL_PATH = "names.xlsx"
OUTPUT_PATH = "labels_print.docx"
ROWS_PER_PAGE = 6
GRID_ROWS = 2
GRID_COLS = 3
FONT_NAME = "Malgun Gothic"
LOGO_PATH = "New Logo.png"
LOGO_WIDTH_MM = 24
PAGE_WIDTH_MM = 297
PAGE_HEIGHT_MM = 210
MARGIN_LEFT_MM = 7
MARGIN_RIGHT_MM = 7
MARGIN_TOP_MM = 4.7
MARGIN_BOTTOM_MM = 4.7
CELL_PADDING_TOP_MM = 2.0
CELL_PADDING_LEFT_MM = 2.8
CELL_PADDING_RIGHT_MM = 2.8
CELL_PADDING_BOTTOM_MM = 2.0
LABEL_WIDTH_MM = 92.5
LABEL_HEIGHT_MM = 88.0


def load_people():
    wb = load_workbook(EXCEL_PATH, data_only=True)
    ws = wb.active
    people = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        name, affiliation = row[:2]
        if not name:
            continue
        people.append(
            {
                "name": str(name).strip(),
                "affiliation": str(affiliation or "").strip(),
            }
        )
    return people


def set_run_font(run, size_pt, bold=False):
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), FONT_NAME)
    r_fonts.set(qn("w:eastAsia"), FONT_NAME)
    r_fonts.set(qn("w:cs"), FONT_NAME)


def name_font_size(text):
    length = len(text)
    if length <= 6:
        return 27
    if length <= 10:
        return 24
    if length <= 16:
        return 21
    if length <= 24:
        return 18
    return 16


def affiliation_font_size(text):
    length = len(text)
    if length <= 10:
        return 16
    if length <= 18:
        return 14.5
    if length <= 28:
        return 12.5
    if length <= 40:
        return 11
    return 10


def clear_cell(cell):
    cell.text = ""
    for paragraph in list(cell.paragraphs):
        p = paragraph._element
        p.getparent().remove(p)
    for table in list(cell.tables):
        tbl = table._tbl
        tbl.getparent().remove(tbl)


def clear_document_body(doc):
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_landscape(section):
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(PAGE_WIDTH_MM)
    section.page_height = Mm(PAGE_HEIGHT_MM)
    section.left_margin = Mm(MARGIN_LEFT_MM)
    section.right_margin = Mm(MARGIN_RIGHT_MM)
    section.top_margin = Mm(MARGIN_TOP_MM)
    section.bottom_margin = Mm(MARGIN_BOTTOM_MM)


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for key, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def hide_cell_borders(cell):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "nil")


def set_row_height_exact(row, height_mm):
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = tr_pr.find(qn("w:trHeight"))
    if tr_height is None:
        tr_height = OxmlElement("w:trHeight")
        tr_pr.append(tr_height)
    tr_height.set(qn("w:val"), str(int(Mm(height_mm).twips)))
    tr_height.set(qn("w:hRule"), "exact")


def make_table_fixed(table):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for col in table.columns:
        col.width = Mm(LABEL_WIDTH_MM)
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        row.height = Mm(LABEL_HEIGHT_MM)
        set_row_height_exact(row, LABEL_HEIGHT_MM)
        for cell in row.cells:
            cell.width = Mm(LABEL_WIDTH_MM)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            hide_cell_borders(cell)
            set_cell_margins(
                cell,
                top=Mm(CELL_PADDING_TOP_MM).twips,
                start=Mm(CELL_PADDING_LEFT_MM).twips,
                bottom=Mm(CELL_PADDING_BOTTOM_MM).twips,
                end=Mm(CELL_PADDING_RIGHT_MM).twips,
            )


def populate_cell(cell, person):
    clear_cell(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    logo_paragraph = cell.add_paragraph()
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_paragraph.paragraph_format.space_before = Pt(0)
    logo_paragraph.paragraph_format.space_after = Pt(7)
    logo_run = logo_paragraph.add_run()
    logo_run.add_picture(LOGO_PATH, width=Mm(LOGO_WIDTH_MM))

    name_paragraph = cell.add_paragraph()
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_paragraph.paragraph_format.space_before = Pt(0)
    name_paragraph.paragraph_format.space_after = Pt(4)
    name_run = name_paragraph.add_run(person["name"])
    set_run_font(name_run, name_font_size(person["name"]), bold=True)

    affiliation_paragraph = cell.add_paragraph()
    affiliation_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affiliation_paragraph.paragraph_format.space_before = Pt(0)
    affiliation_paragraph.paragraph_format.space_after = Pt(0)
    affiliation_run = affiliation_paragraph.add_run(person["affiliation"])
    set_run_font(
        affiliation_run,
        affiliation_font_size(person["affiliation"]),
        bold=False,
    )

def build_page_table(doc):
    table = doc.add_table(rows=GRID_ROWS, cols=GRID_COLS)
    make_table_fixed(table)
    return table


def main():
    people = load_people()
    doc = Document(TEMPLATE_PATH)
    section = doc.sections[0]
    set_landscape(section)
    clear_document_body(doc)
    pages = [people[i : i + ROWS_PER_PAGE] for i in range(0, len(people), ROWS_PER_PAGE)]

    for page_index, batch in enumerate(pages):
        if page_index > 0:
            doc.add_page_break()
        table = build_page_table(doc)

        cells = [cell for row in table.rows for cell in row.cells]
        for idx, cell in enumerate(cells):
            if idx < len(batch):
                populate_cell(cell, batch[idx])
            else:
                clear_cell(cell)

    doc.save(OUTPUT_PATH)
    print(f"Created {OUTPUT_PATH} with {len(people)} labels across {len(pages)} pages.")


if __name__ == "__main__":
    main()
